#!/usr/bin/env python3
"""Render tailored client documents from content/policies/*.md + an intake JSON.

Usage:
  engine/render.py --intake content/intake.example.json --out build/example [--only slug1,slug2] [--html]

Outputs one DOCX per document (and optional HTML), plus INDEX.md listing documents by standard.
"""
import argparse, datetime, json, os, re, sys
import yaml
from jinja2 import Environment, StrictUndefined, TemplateError
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
CONTENT = os.path.join(ROOT, "content", "policies")
FRONT = re.compile(r"^---\s*\n(.*?)\n---\s*\n", re.S)

# ---------- Jinja ----------
def fmt_date(v):
    if not v:
        return "[TO CONFIRM]"
    try:
        d = datetime.date.fromisoformat(str(v)[:10])
        return f"{d.day} {d.strftime('%B %Y')}"
    except ValueError:
        return str(v)

def make_env():
    env = Environment(undefined=StrictUndefined, trim_blocks=True, lstrip_blocks=True, autoescape=False)
    env.filters["date"] = fmt_date
    return env

def load_docs(only=None):
    docs = []
    for fn in sorted(os.listdir(CONTENT)):
        if not fn.endswith(".md"):
            continue
        slug = fn[:-3]
        if only and slug not in only:
            continue
        raw = open(os.path.join(CONTENT, fn), encoding="utf-8").read()
        m = FRONT.match(raw)
        if not m:
            raise SystemExit(f"{fn}: missing YAML front matter")
        meta = yaml.safe_load(m.group(1)) or {}
        meta.setdefault("slug", slug)
        docs.append((meta, raw[m.end():]))
    return docs

# ---------- Markdown → DOCX ----------
INLINE = re.compile(r"(\*\*.+?\*\*)")

def add_inline(par, text):
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            par.add_run(part[2:-2]).bold = True
        else:
            par.add_run(part)

def style_doc(doc, org_name, title):
    st = doc.styles["Normal"]
    st.font.name = "Calibri"; st.font.size = Pt(11)
    for s in doc.sections:
        s.left_margin = s.right_margin = Cm(2.2); s.top_margin = s.bottom_margin = Cm(2)
        hdr = s.header.paragraphs[0]; hdr.text = f"{org_name} — {title}"; hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for r in hdr.runs: r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        ftr = s.footer.paragraphs[0]; ftr.text = "Controlled document — uncontrolled when printed"; ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in ftr.runs: r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

def md_to_docx(md, org_name, title):
    doc = Document(); style_doc(doc, org_name, title)
    lines = md.splitlines(); i = 0; num_counter = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            num_counter = 0; i += 1; continue
        if line.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
                    rows.append(cells)
                i += 1
            if rows:
                ncol = max(len(r) for r in rows)
                t = doc.add_table(rows=0, cols=ncol); t.style = "Table Grid"
                for ri, r in enumerate(rows):
                    cells = t.add_row().cells
                    for ci in range(ncol):
                        cells[ci].text = ""
                        p = cells[ci].paragraphs[0]; add_inline(p, r[ci] if ci < len(r) else "")
                        for run in p.runs:
                            run.font.size = Pt(9.5)
                            if ri == 0: run.bold = True
                doc.add_paragraph()
            continue
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            level = len(m.group(1)); text = m.group(2).strip()
            if level == 1:
                doc.add_heading(text, 0)
            else:
                doc.add_heading(text, min(level - 1, 3))
            i += 1; continue
        m = re.match(r"^\s*[-*]\s+(.*)", line)
        if m:
            p = doc.add_paragraph(style="List Bullet"); add_inline(p, m.group(1)); i += 1; continue
        m = re.match(r"^\s*\d+[.)]\s+(.*)", line)
        if m:
            p = doc.add_paragraph(style="List Number"); add_inline(p, m.group(1)); i += 1; continue
        # paragraph: merge consecutive plain lines
        buf = [line.strip()]; i += 1
        while i < len(lines) and lines[i].strip() and not re.match(r"^(#{1,4}\s|\s*[-*]\s|\s*\d+[.)]\s|\|)", lines[i]):
            buf.append(lines[i].strip()); i += 1
        p = doc.add_paragraph(); add_inline(p, " ".join(buf))
    return doc

def md_to_html(md, title):
    import html as h
    out = [f"<!doctype html><meta charset='utf-8'><title>{h.escape(title)}</title><style>body{{font-family:Calibri,Arial,sans-serif;max-width:860px;margin:2rem auto;line-height:1.5}}table{{border-collapse:collapse}}td,th{{border:1px solid #999;padding:4px 8px;font-size:.9em}}</style>"]
    in_list = None
    for line in md.splitlines():
        if line.startswith("|"):
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            if all(re.fullmatch(r":?-{2,}:?", c) for c in cells): continue
            out.append("<tr>" + "".join(f"<td>{h.escape(c)}</td>" for c in cells) + "</tr>")
            continue
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m: out.append(f"<h{len(m.group(1))}>{h.escape(m.group(2))}</h{len(m.group(1))}>"); continue
        m = re.match(r"^\s*[-*]\s+(.*)", line)
        if m: out.append(f"<li>{h.escape(m.group(1))}</li>"); continue
        if line.strip(): out.append(f"<p>{h.escape(line.strip())}</p>")
    html_ = "\n".join(out)
    html_ = re.sub(r"(<tr>.*?</tr>\n?)+", lambda m: "<table>" + m.group(0) + "</table>", html_, flags=re.S)
    html_ = re.sub(r"(<li>.*?</li>\n?)+", lambda m: "<ul>" + m.group(0) + "</ul>", html_, flags=re.S)
    return html_.replace("**", "")

# ---------- main ----------
def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--intake", required=True); ap.add_argument("--out", required=True)
    ap.add_argument("--only"); ap.add_argument("--html", action="store_true"); ap.add_argument("--check", action="store_true", help="render only, write nothing")
    a = ap.parse_args()
    intake = json.load(open(a.intake, encoding="utf-8"))
    intake.setdefault("meta", {}).setdefault("generated_on", datetime.date.today().isoformat())
    ctx = {"intake": intake, "org": intake["org"], "today": datetime.date.today().isoformat()}
    env = make_env(); only = set(a.only.split(",")) if a.only else None
    os.makedirs(a.out, exist_ok=True) if not a.check else None
    index = {}; errors = 0; produced = 0
    for meta, body in load_docs(only):
        slug = meta["slug"]
        try:
            applies = env.from_string("{{ (" + str(meta.get("applies_if", "true")) + ") }}").render(**ctx).strip().lower() == "true"
        except TemplateError as e:
            print(f"ERROR {slug}: applies_if: {e}"); errors += 1; continue
        if not applies:
            print(f"skip  {slug} (applies_if false)"); continue
        try:
            text = env.from_string(body).render(**ctx)
        except TemplateError as e:
            print(f"ERROR {slug}: {e}"); errors += 1; continue
        leftovers = re.findall(r"\{[{%].*?[}%]\}", text)
        if leftovers:
            print(f"ERROR {slug}: unrendered tags {leftovers[:3]}"); errors += 1; continue
        title = meta.get("title", slug)
        if not a.check:
            md_to_docx(text, intake["org"]["name"], title).save(os.path.join(a.out, f"{slug}.docx"))
            if a.html:
                open(os.path.join(a.out, f"{slug}.html"), "w", encoding="utf-8").write(md_to_html(text, title))
        for code in meta.get("standards", []):
            index.setdefault(code, []).append((slug, title))
        produced += 1
        print(f"ok    {slug} ({len(text.split())} words, {meta.get('doc_type')})")
    if not a.check:
        with open(os.path.join(a.out, "INDEX.md"), "w", encoding="utf-8") as f:
            f.write(f"# Evidence index — {intake['org']['name']}\n\nGenerated {ctx['today']}. Documents listed under each NDIS Practice Standards outcome they evidence.\n\n")
            for code in sorted(index):
                f.write(f"## {code}\n" + "".join(f"- {t} (`{s}.docx`)\n" for s, t in index[code]) + "\n")
    print(f"\n{produced} documents rendered, {errors} errors")
    sys.exit(1 if errors else 0)

if __name__ == "__main__":
    main()
